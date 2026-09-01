from __future__ import annotations

import csv
import hashlib
import re
import sys
from pathlib import Path

from docx import Document
from pypdf import PdfReader


RESEARCH_DIR = Path(__file__).resolve().parent
INVENTORY = RESEARCH_DIR / "inventory.csv"
CORPUS_DIR = RESEARCH_DIR / "corpus"
MANIFEST = RESEARCH_DIR / "corpus-manifest.csv"
MANIFEST_FIELDS = [
    "Course",
    "Role",
    "Extension",
    "Status",
    "Units",
    "Characters",
    "SHA256",
    "CorpusFile",
    "RelativePath",
    "FullPath",
]

EXCLUDED_TOP_LEVEL = {"题本"}
EXCLUDED_FILENAMES = {
    "数量关系理论课讲义224.pdf",
    "政治常识一本通.pdf",
}


def classify(relative_path: str) -> str:
    name = Path(relative_path).name
    if "精读" in name or "全文勾画" in name:
        return "精读/勾画"
    if "批改标准" in name:
        return "批改标准"
    if "解析" in name or "参考答案" in name or "做题痕迹" in name:
        return "解析/参考答案"
    if "板书" in name or "笔记" in name:
        return "板书/课堂笔记"
    if "范文" in name:
        return "范文"
    if "作业" in name:
        return "课后作业"
    if any(token in name for token in ("题本", "用题", "真题", "预习")):
        return "题本/预习材料"
    if any(token in name for token in ("报告", "背诵", "背景", "常识", "时政")):
        return "背景素材"
    return "讲义/其他"


def course_name(top: str, relative_path: str) -> str:
    if top in {
        "01.上课用题",
        "03.课后作业",
    } or re.fullmatch(r"\d{4}第.+节", top):
        return "早期破题阵/扎马步课"
    if top in {"大作文范文", "大作文课题本"}:
        return "大作文专题（散装）"
    if top == "讲义":
        if "\\01" in relative_path or "\\02" in relative_path:
            return "2026省考大决战/国考点将台（混合目录）"
        return "2026省考大决战/国考点将台（混合目录）"
    if top == "(root)":
        return "早期破题阵/扎马步课（根目录散件）"
    return top


def run_text(run) -> str:
    text = run.text
    if not text:
        return ""
    tags: list[str] = []
    if run.bold:
        tags.append("粗体")
    if run.underline:
        tags.append("下划线")
    try:
        highlight = run.font.highlight_color
    except ValueError:
        # w:highlight="none" is a valid Word value but unsupported by the
        # python-docx enum used by the bundled runtime.
        highlight = None
    if highlight is not None:
        # Some course files explicitly store w:highlight="none".  python-docx
        # can parse that enum, but its __str__ implementation raises because
        # Word has no display-name mapping for the value.
        tags.append(f"高亮索引:{int(highlight)}")
    if run.font.color is not None and run.font.color.rgb is not None:
        tags.append(f"色:{run.font.color.rgb}")
    if not tags:
        return text
    return f"〔{'|'.join(tags)}〕{text}〔/〕"


def extract_docx(path: Path) -> tuple[str, int]:
    doc = Document(path)
    blocks: list[str] = []
    for paragraph in doc.paragraphs:
        value = "".join(run_text(run) for run in paragraph.runs).strip()
        if value:
            blocks.append(value)
    for table_index, table in enumerate(doc.tables, start=1):
        blocks.append(f"[表格 {table_index}]")
        for row in table.rows:
            cells = []
            for cell in row.cells:
                value = " / ".join(
                    "".join(run_text(run) for run in p.runs).strip()
                    for p in cell.paragraphs
                    if p.text.strip()
                )
                cells.append(value)
            blocks.append(" | ".join(cells))
    return "\n".join(blocks), len(blocks)


def extract_pdf(path: Path) -> tuple[str, int]:
    reader = PdfReader(path, strict=False)
    pages: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            value = page.extract_text() or ""
        except Exception as exc:  # retain partial documents and record the page failure
            value = f"[第 {page_index} 页提取失败：{type(exc).__name__}: {exc}]"
        pages.append(f"\n[第 {page_index} 页]\n{value.strip()}")
    return "\n".join(pages).strip(), len(reader.pages)


def safe_name(relative_path: str, digest: str) -> str:
    stem = Path(relative_path).stem
    stem = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_-]+", "_", stem).strip("_")
    return f"{digest[:12]}__{stem[:70]}.txt"


def main() -> int:
    CORPUS_DIR.mkdir(parents=True, exist_ok=True)
    with INVENTORY.open("r", encoding="utf-8-sig", newline="") as handle:
        source_rows = list(csv.DictReader(handle))

    previous_by_path: dict[str, dict[str, str]] = {}
    if MANIFEST.exists():
        with MANIFEST.open("r", encoding="utf-8-sig", newline="") as handle:
            previous_by_path = {
                row["RelativePath"]: row for row in csv.DictReader(handle)
            }

    output_rows: list[dict[str, str | int]] = []
    for index, row in enumerate(source_rows, start=1):
        top = row["Top"]
        relative_path = row["RelativePath"]
        extension = row["Extension"].lower()
        source = Path(row["FullPath"])
        if top in EXCLUDED_TOP_LEVEL or source.name in EXCLUDED_FILENAMES:
            continue
        if extension not in {".docx", ".pdf"}:
            output_rows.append(
                {
                    "Course": course_name(top, relative_path),
                    "Role": classify(relative_path),
                    "Extension": extension,
                    "Status": "待转换" if extension in {".doc", ".wps"} else "暂不抽取",
                    "Units": "",
                    "Characters": "",
                    "SHA256": "",
                    "CorpusFile": "",
                    "RelativePath": relative_path,
                    "FullPath": str(source),
                }
            )
            continue

        digest = hashlib.sha256(source.read_bytes()).hexdigest().upper()
        target = CORPUS_DIR / safe_name(relative_path, digest)
        previous = previous_by_path.get(relative_path)
        if previous and previous.get("Status") == "已抽取" and target.exists():
            output_rows.append(previous)
            continue
        try:
            if extension == ".docx":
                text, units = extract_docx(source)
            else:
                text, units = extract_pdf(source)
            status = "已抽取" if text.strip() else "无可提取文本/OCR待办"
            header = (
                f"来源：{relative_path}\n"
                f"课程：{course_name(top, relative_path)}\n"
                f"资料角色：{classify(relative_path)}\n"
                f"SHA256：{digest}\n"
                "说明：〔色/高亮/粗体〕标记来自 DOCX 原始文字格式。\n\n"
            )
            target.write_text(header + text, encoding="utf-8")
            output_rows.append(
                {
                    "Course": course_name(top, relative_path),
                    "Role": classify(relative_path),
                    "Extension": extension,
                    "Status": status,
                    "Units": units,
                    "Characters": len(text),
                    "SHA256": digest,
                    "CorpusFile": str(target.relative_to(RESEARCH_DIR)),
                    "RelativePath": relative_path,
                    "FullPath": str(source),
                }
            )
        except Exception as exc:
            output_rows.append(
                {
                    "Course": course_name(top, relative_path),
                    "Role": classify(relative_path),
                    "Extension": extension,
                    "Status": f"失败:{type(exc).__name__}:{exc}",
                    "Units": "",
                    "Characters": "",
                    "SHA256": digest,
                    "CorpusFile": "",
                    "RelativePath": relative_path,
                    "FullPath": str(source),
                }
            )
        if index % 25 == 0:
            print(f"processed {index}/{len(source_rows)}", file=sys.stderr)

    with MANIFEST.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=MANIFEST_FIELDS)
        writer.writeheader()
        writer.writerows(output_rows)

    extracted = sum(row["Status"] == "已抽取" for row in output_rows)
    pending = sum(row["Status"] == "待转换" for row in output_rows)
    failed = sum(str(row["Status"]).startswith("失败:") for row in output_rows)
    print(f"extracted={extracted} pending_legacy={pending} failed={failed}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
